"""
按导入模版.docx 格式生成章末习题和期末考试 .docx 文件。

题库全部基于真实的计算机硬件知识，与课程内容一一对应。

模板格式:
  - 一、单项选择题 (5 题) — 题干 + A/B/C/D 选项 + 答案标注 + [解析]
  - 二、多项选择题 (5 题) — 同上
  - 三、判断题 (5 题) — 题干 + 答案 (√/×) + [解析]

期末考试:
  - 单选20 + 多选10 + 填空5 + 判断15 = 50 题，每题 2 分 = 100 分

用法:
  from scrape_new.upload.exercise_docx import generate_exercise_docx, generate_final_exam_docx
  generate_exercise_docx(chapter_data, output_path)
  generate_final_exam_docx(outline_data, output_path)
"""

from __future__ import annotations

import re
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── 常量 ──────────────────────────────────────────────────────

FONT_NAME_CN = "宋体"
FONT_NAME_EN = "Times New Roman"
FONT_SIZE = Pt(12)
TITLE_SIZE = Pt(16)


# ═══════════════════════════════════════════════════════════════
#  题库 — 根据每章实际教学内容编写
# ═══════════════════════════════════════════════════════════════

# BANK 已迁移到 exercise_bank.json,由 _get_bank() 加载

def _load_bank() -> dict:
    """从 exercise_bank.json 加载题库(独立于 Python 源码,方便换课)。"""
    _p = Path(__file__).parent / "exercise_bank.json"
    if not _p.exists():
        raise FileNotFoundError(f"题库文件不存在: {_p}")
    with open(_p, "r", encoding="utf-8") as _f:
        return json.loads(_f.read())

# 延迟加载:首次访问时从 JSON 文件读取,后续调用复用缓存
_BANK_CACHE: dict | None = None

def _get_bank() -> dict:
    global _BANK_CACHE
    if _BANK_CACHE is None:
        raw = _load_bank()
        _BANK_CACHE = {int(k): v for k, v in raw.items()}  # JSON key 是 str → 转 int
    return _BANK_CACHE


# ── 期末考试题库 ──
FINAL_EXAM_SINGLE = [
    # 20 题单选，覆盖全部章节
    {"s": "世界上第一台电子数字计算机 ENIAC 诞生于哪一年", "a": "B",
     "o": ["A. 1945 年", "B. 1946 年", "C. 1947 年", "D. 1950 年"],
     "r": "ENIAC 于 1946 年在美国宾夕法尼亚大学正式投入使用。"},
    {"s": "冯·诺依曼体系结构的五大部件中不包括", "a": "D",
     "o": ["A. 运算器", "B. 控制器", "C. 存储器", "D. 显卡"],
     "r": "冯·诺依曼体系五大部件为运算器、控制器、存储器、输入设备和输出设备。"},
    {"s": "目前个人计算机 CPU 市场的主要两家厂商是", "a": "B",
     "o": ["A. 三星和 LG", "B. Intel 和 AMD", "C. NVIDIA 和 AMD", "D. Intel 和 NVIDIA"],
     "r": "Intel 和 AMD 是全球 PC CPU 市场的两大主要厂商。"},
    {"s": "CPU 缓存中速度最快、容量最小的是", "a": "A",
     "o": ["A. L1 Cache", "B. L2 Cache", "C. L3 Cache", "D. 硬盘缓存"],
     "r": "L1 缓存集成在 CPU 核心内部，速度最快但容量最小（通常几十KB）。"},
    {"s": "主板上纽扣电池的作用是", "a": "B",
     "o": ["A. 为主板供电", "B. 保持 CMOS 中的 BIOS 设置和系统时间", "C. 为 CPU 供电", "D. 为内存供电"],
     "r": "CMOS 电池用于关机后维持 BIOS 设置和时间，取下后设置会恢复默认。"},
    {"s": "目前主流计算机使用的内存类型是", "a": "C",
     "o": ["A. SDRAM", "B. DDR3", "C. DDR4/DDR5", "D. RDRAM"],
     "r": "当前主流平台使用 DDR4（上一代主流）和 DDR5（最新一代）。"},
    {"s": "显卡的核心芯片称为", "a": "B",
     "o": ["A. CPU", "B. GPU（图形处理器）", "C. RAM", "D. NPU"],
     "r": "GPU（Graphics Processing Unit）是显卡的核心，负责图形渲染和并行计算。"},
    {"s": "显示器分辨率 1920×1080 通常被称为", "a": "B",
     "o": ["A. 720p", "B. 1080p（全高清/FHD）", "C. 2K", "D. 4K"],
     "r": "1920×1080 俗称 1080p 或全高清（Full HD）。"},
    {"s": "以下哪种硬盘速度最快", "a": "D",
     "o": ["A. 5400 转机械硬盘", "B. 7200 转机械硬盘", "C. SATA 固态硬盘", "D. NVMe M.2 固态硬盘"],
     "r": "NVMe M.2 SSD 利用 PCIe 通道，顺序读取可达 7000MB/s+。"},
    {"s": "80 PLUS 电源认证中，转换效率最高的是", "a": "D",
     "o": ["A. 铜牌", "B. 银牌", "C. 金牌", "D. 钛金牌"],
     "r": "80 PLUS 认证从低到高：白牌→铜牌→银牌→金牌→白金牌→钛金牌。"},
    {"s": "安装 CPU 时正确的做法是", "a": "C",
     "o": ["A. 用力按压四角", "B. 使用螺丝刀辅助", "C. 对准三角形标记，自然落入插槽", "D. 先涂抹硅脂再放入"],
     "r": "CPU 应对准防呆标志自然落入插槽，不需要外力按压。硅脂应在装好 CPU 后涂在顶盖上。"},
    {"s": "制作 Windows 启动 U 盘时 U 盘文件系统通常为", "a": "B",
     "o": ["A. NTFS", "B. FAT32", "C. exFAT", "D. ext4"],
     "r": "UEFI 需要 FAT32 格式的引导分区。部分工具（如 Rufus）可自动处理。"},
    {"s": "MBR 分区表最多支持几个主分区", "a": "C",
     "o": ["A. 2 个", "B. 3 个", "C. 4 个", "D. 不限"],
     "r": "MBR 分区表最多支持 4 个主分区（或 3 主 + 1 扩展）。GPT 支持 128 个分区。"},
    {"s": "Windows 安装完成后的第一步应该", "a": "B",
     "o": ["A. 立刻装游戏", "B. 安装主板驱动和显卡驱动", "C. 删除安装 U 盘文件", "D. 格式化所有分区"],
     "r": "系统安装后首先应安装硬件驱动，确保各部件正常工作。"},
    {"s": "进入 BIOS 设置的常见按键是", "a": "B",
     "o": ["A. Ctrl+Alt+Del", "B. Del 或 F2", "C. Windows 键", "D. Alt+F4"],
     "r": "大部分主板使用 Del 或 F2 进入 BIOS/UEFI 设置界面。"},
    {"s": "误删除文件后第一步应该", "a": "A",
     "o": ["A. 停止对硬盘的所有写入操作", "B. 重新安装操作系统", "C. 格式化硬盘", "D. 反复重启"],
     "r": "文件删除后数据仍在磁盘上，停止写入可防止数据被覆盖，提高恢复成功率。"},
    {"s": "SSD 删除文件后恢复难度大的主要原因是", "a": "B",
     "o": ["A. SSD 速度太快", "B. TRIM 功能主动擦除已删文件数据", "C. SSD 容量太小", "D. SSD 使用不同文件系统"],
     "r": "TRIM 指令通知 SSD 控制器擦除无效页面，提升写入效率的同时使被删数据难以恢复。"},
    {"s": "计算机维修中'最小系统法'的核心思路是", "a": "C",
     "o": ["A. 使用最小功率电源", "B. 只运行一个程序", "C. 保留最少核心部件通电测试，逐步添加定位故障", "D. 使用安全模式"],
     "r": "最小系统法只保留 CPU+主板+内存+电源（+显卡如无核显），确认正常后逐步加回其他部件。"},
    {"s": "1KB 等于多少字节", "a": "C",
     "o": ["A. 1000 字节", "B. 1000 位", "C. 1024 字节", "D. 1024 位"],
     "r": "计算机采用二进制，1KB = 1024B。"},
    {"s": "开启 XMP/DOCP 的主要目的是", "a": "A",
     "o": ["A. 让内存运行在标称高频下", "B. 提高显卡性能", "C. 延长硬盘寿命", "D. 降低 CPU 温度"],
     "r": "XMP（Intel）/DOCP（AMD）开启后内存以厂商标称频率运行，否则默认低频。"},
]

FINAL_EXAM_MULTI = [
    {"s": "微型计算机的主要特点包括", "a": "ABCD",
     "o": ["A. 体积小、重量轻", "B. 价格低廉", "C. 可靠性高", "D. 环境适应性强"],
     "r": "微型计算机具有体积小、价格低、可靠性高、使用方便等特点。"},
    {"s": "影响 CPU 性能的主要因素有", "a": "ABCD",
     "o": ["A. 核心数量", "B. 主频高低", "C. 缓存大小", "D. 制程工艺"],
     "r": "以上四项均直接影响 CPU 性能。"},
    {"s": "主板的主要功能包括", "a": "ABCD",
     "o": ["A. 连接和协调各硬件", "B. 提供电源分配", "C. 板载各类接口", "D. 承载核心部件"],
     "r": "主板是计算机各部件连接枢纽，以上均为核心功能。"},
    {"s": "选购内存时需要考虑的参数有", "a": "ABCD",
     "o": ["A. 容量", "B. 频率", "C. DDR 代数", "D. 时序（CL 值）"],
     "r": "容量、频率、DDR 代数和时序是内存的四大核心参数。"},
    {"s": "以下属于显卡输出接口的有", "a": "ABD",
     "o": ["A. HDMI", "B. DisplayPort", "C. RJ45", "D. DVI"],
     "r": "HDMI、DP 和 DVI 是显卡的视频输出接口。RJ45 是网络接口。"},
    {"s": "选购显示器时需要考虑的参数有", "a": "ABCD",
     "o": ["A. 分辨率", "B. 刷新率", "C. 面板类型", "D. 色域"],
     "r": "以上均为显示器选购的核心参数。"},
    {"s": "相比机械硬盘，固态硬盘的优势有", "a": "ABCD",
     "o": ["A. 读写速度更快", "B. 抗震性能更好", "C. 功耗更低", "D. 无噪音"],
     "r": "SSD 无机械部件，在速度、抗震、功耗和噪音方面全面优于 HDD。"},
    {"s": "组装计算机时的注意事项有", "a": "ABCD",
     "o": ["A. 防静电", "B. 阅读主板说明书", "C. 检查部件外观", "D. 断电操作"],
     "r": "以上均为装机基本安全规范。"},
    {"s": "常见的分区表类型有", "a": "AB",
     "o": ["A. MBR", "B. GPT", "C. NTFS", "D. FAT32"],
     "r": "MBR 和 GPT 是分区表。NTFS/FAT32 是文件系统。"},
    {"s": "数据恢复的基本原则包括", "a": "ABD",
     "o": ["A. 停止对目标盘的写入", "B. 恢复软件不要装到待恢复盘", "C. 反复重启电脑", "D. 恢复出的文件存到其他盘"],
     "r": "反复重启不会提高恢复成功率，反而可能在启动过程中写入数据覆盖待恢复文件。"},
]

# 解析器不支持填空题(___格式识别为 FillBlank 但不提取答案,TemplateID=None)
# 用 5 道单选题替代,总分仍是 100（25单选=50分 + 10多选=20分 + 15判断=30分）
FINAL_EXAM_SINGLE_MORE = [
    {"s": "计算机系统的两大组成部分是硬件系统和", "a": "A",
     "o": ["A. 软件系统", "B. 电源系统", "C. 网络系统", "D. 散热系统"],
     "r": "计算机系统由硬件系统（物理设备）和软件系统（程序和数据）两部分组成。"},
    {"s": "CPU 内部负责算术运算和逻辑判断的部件是", "a": "B",
     "o": ["A. 控制器（CU）", "B. 运算器（ALU）", "C. 寄存器组", "D. 缓存"],
     "r": "运算器（ALU）是 CPU 内部执行算术运算和逻辑运算的核心部件。"},
    {"s": "固态硬盘（SSD）使用的存储介质是", "a": "C",
     "o": ["A. 磁性盘片", "B. 光盘介质", "C. NAND Flash 闪存芯片", "D. DRAM 芯片"],
     "r": "SSD 基于 NAND Flash 闪存芯片存取数据，具有读写快、抗震等优势。"},
    {"s": "以下哪个是常用的免费系统启动 U 盘制作工具", "a": "D",
     "o": ["A. Photoshop", "B. WinRAR", "C. Notepad++", "D. Rufus"],
     "r": "Rufus 和 Ventoy 均为常用的免费启动盘制作工具。"},
    {"s": "计算机故障排查的基本原则是", "a": "A",
     "o": ["A. 先软件后硬件，先简单后复杂", "B. 先硬件后软件，先复杂后简单", "C. 直接替换所有硬件", "D. 先格式化重装系统"],
     "r": "故障排查遵循先简后繁、先软后硬、先外后内的原则。"},
]

FINAL_EXAM_JUDGE = [
    {"s": "1GB 等于 1024MB。（√）"},
    {"s": "Intel 和 AMD 的 CPU 接口可以互换使用。（×）",
     "r": "Intel（LGA）和 AMD（PGA/LGA）使用不同的插槽标准，物理不兼容。"},
    {"s": "主板上的纽扣电池没电会导致 BIOS 设置丢失。（√）"},
    {"s": "DDR5 内存可以插入 DDR4 主板使用。（×）",
     "r": "DDR5 和 DDR4 的物理接口（针脚数和缺口位置）不同，不能互换。"},
    {"s": "安装双通道内存时两条内存应插在同色插槽（如 A2+B2）。（√）"},
    {"s": "GPU 是显卡的核心芯片，负责图形渲染。（√）"},
    {"s": "刷新率 144Hz 的显示器每秒最多显示 144 帧。（√）"},
    {"s": "NVMe M.2 固态硬盘的速度比 SATA 固态硬盘快。（√）"},
    {"s": "80 PLUS 认证等级越高，电源转换效率越高。（√）"},
    {"s": "组装计算机时应在通电状态下插拔硬件以测试是否正常。（×）",
     "r": "通电状态下插拔硬件可能损坏硬件或造成人身伤害，必须断电操作。"},
    {"s": "制作启动 U 盘会格式化 U 盘，原有数据会丢失。（√）"},
    {"s": "GPT 分区表支持大于 2TB 的硬盘。（√）"},
    {"s": "安装操作系统前不需要备份数据。（×）",
     "r": "安装系统可能覆盖目标分区，安装前务必备份重要文件。"},
    {"s": "清除主板的 CMOS 可以让 BIOS 设置恢复默认。（√）"},
    {"s": "SSD 的 TRIM 功能会使被删除文件更难恢复。（√）"},
]


# ═══════════════════════════════════════════════════════════════
#  docx 生成函数
# ═══════════════════════════════════════════════════════════════

def _create_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME_EN
    style.font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    return doc


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = False
    run.font.size = FONT_SIZE
    run.font.name = FONT_NAME_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)


def _add_para(doc: Document, text: str, bold: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = FONT_SIZE
    run.font.name = FONT_NAME_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    if bold:
        run.bold = True


def _add_title(doc: Document, name: str, is_final: bool = False) -> None:
    title = f"{name} — 期末考试" if is_final else f"{name} — 章末测试"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = TITLE_SIZE
    run.font.name = FONT_NAME_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    doc.add_paragraph()


# ─── 章末习题 ────────────────────────────────────────────────

def generate_exercise_docx(chapter_data: dict[str, Any], output_path: Path) -> Path:
    ch_index = chapter_data.get("ch_index", 1)
    ch_name = chapter_data.get("ch_name", f"第{ch_index}章")
    # 去掉"第X章"前缀
    clean = re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", ch_name)

    bank = _get_bank().get(ch_index)
    if not bank:
        raise ValueError(f"第{ch_index}章 题库未编写（BANK 中无此章节）")

    single, multi, judge = bank[0], bank[1], bank[2]

    doc = _create_document()
    _add_title(doc, clean)

    # 一、单选
    _add_section_heading(doc, "一、单项选择题")
    for i, q in enumerate(single, 1):
        _add_para(doc, f"{i}. {q['s']}（{q['a']}）")
        for o in q["o"]:
            _add_para(doc, o)
        _add_para(doc, f"[解析]：{q['r']}")
        doc.add_paragraph()

    # 二、多选
    _add_section_heading(doc, "二、多项选择题")
    for i, q in enumerate(multi, 1):
        _add_para(doc, f"{i}. {q['s']}（{q['a']}）")
        for o in q["o"]:
            _add_para(doc, o)
        _add_para(doc, f"[解析]：{q['r']}")
        doc.add_paragraph()

    # 三、判断
    _add_section_heading(doc, "三、判断题")
    for i, q in enumerate(judge, 1):
        s = q["s"]
        remark = q.get("r", "")
        _add_para(doc, f"{i}. {s}")
        if remark:
            _add_para(doc, f"[解析]：{remark}")
        doc.add_paragraph()

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  [习题] 已生成: {output_path}")
    return output_path


# ─── 期末考试 ────────────────────────────────────────────────

def generate_final_exam_docx(outline_data: dict[str, Any], output_path: Path) -> Path:
    course_title = outline_data.get("course_title", "本课程")
    clean_title = re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", course_title)

    doc = _create_document()
    _add_title(doc, clean_title, is_final=True)

    # 一、单选 25（20 + 5 extra 替代填空题）
    _add_section_heading(doc, "一、单项选择题")
    _add_para(doc, "（每题 2 分，共 50 分）")
    for i, q in enumerate(FINAL_EXAM_SINGLE, 1):
        _add_para(doc, f"{i}. {q['s']}（{q['a']}）")
        for o in q["o"]:
            _add_para(doc, o)
        _add_para(doc, f"[解析]：{q['r']}")
        doc.add_paragraph()
    for i, q in enumerate(FINAL_EXAM_SINGLE_MORE, 21):
        _add_para(doc, f"{i}. {q['s']}（{q['a']}）")
        for o in q["o"]:
            _add_para(doc, o)
        _add_para(doc, f"[解析]：{q['r']}")
        doc.add_paragraph()

    # 二、多选 10
    _add_section_heading(doc, "二、多项选择题")
    _add_para(doc, "（每题 2 分，共 20 分）")
    for i, q in enumerate(FINAL_EXAM_MULTI, 1):
        _add_para(doc, f"{i}. {q['s']}（{q['a']}）")
        for o in q["o"]:
            _add_para(doc, o)
        _add_para(doc, f"[解析]：{q['r']}")
        doc.add_paragraph()

    # 三、判断 15
    _add_section_heading(doc, "三、判断题")
    _add_para(doc, "（每题 2 分，共 30 分）")
    for i, q in enumerate(FINAL_EXAM_JUDGE, 1):
        s = q["s"]
        remark = q.get("r", "")
        _add_para(doc, f"{i}. {s}")
        if remark:
            _add_para(doc, f"[解析]：{remark}")
        doc.add_paragraph()

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  [期末] 已生成: {output_path}")
    return output_path


# ─── CLI demo ─────────────────────────────────────────────────

if __name__ == "__main__":
    import sys, json
    if len(sys.argv) < 3:
        print("用法: python exercise_docx.py <outline.json> <输出目录>")
        sys.exit(1)
    outline_path = Path(sys.argv[1])
    out_dir = Path(sys.argv[2])
    data = json.loads(outline_path.read_text(encoding="utf-8"))
    chapters = data.get("chapters", [])
    out_dir.mkdir(parents=True, exist_ok=True)
    for ch in chapters:
        idx = ch.get("index", 1)
        ch_data = {
            "ch_index": idx,
            "ch_name": ch.get("title", f"第{idx}章"),
            "lessons": ch.get("lessons", []),
        }
        fname = f"第{idx}章_{ch_data['ch_name']}_习题.docx"
        generate_exercise_docx(ch_data, out_dir / fname)
